from fastapi import FastAPI, Form, Query
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from rag_engine import ask_question
from fpdf import FPDF
from docx import Document
import uuid
import os
import re
import time

app = FastAPI(
    title="Arabic Legal AI Assistant 🇸🇦",
    description="استشارة قانونية ومالية وإدارية ذكية مبنية على القانون السعودي.",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For dev only!
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create temp directory if it doesn't exist
TEMP_DIR = "temp_files"
os.makedirs(TEMP_DIR, exist_ok=True)

# Serve the HTML page directly from FastAPI
@app.get("/", response_class=HTMLResponse)
def read_root():
    return """
<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="UTF-8" />
<title>المساعد القانوني الذكي</title>
<style>
body { font-family: Tahoma; background: #f0f0f0; padding: 2em; }
textarea { width: 100%; height: 150px; font-size: 16px; }
button { padding: 10px 20px; font-size: 16px; }
#answer { margin-top: 20px; background: #fff; padding: 20px; border-radius: 10px; }
.export-links { margin-top: 15px; padding: 15px; background: #e8f5e8; border-radius: 8px; }
.export-links a { 
    display: inline-block; 
    margin: 5px 10px 5px 0; 
    padding: 8px 15px; 
    background: #28a745; 
    color: white; 
    text-decoration: none; 
    border-radius: 5px; 
}
.export-links a:hover { background: #1e7e34; }
</style>
</head>
<body>
<h2>📋 اكتب سؤالك القانوني هنا</h2>
<textarea id="query" placeholder="اكتب سؤالك هنا..."></textarea><br>
<button onclick="sendQuestion()">🔍 أرسل السؤال</button>
<div id="answer"></div>
<div id="export"></div>

<script>
async function sendQuestion() {
    const query = document.getElementById('query').value.trim();
    if (!query) {
        alert('الرجاء كتابة سؤال قبل الإرسال');
        return;
    }

    // Show loading
    document.getElementById('answer').innerHTML = '<p>⏳ جاري معالجة السؤال...</p>';
    document.getElementById('export').innerHTML = '';

    try {
        const formData = new FormData();
        formData.append('query', query);
        
        const response = await fetch('/ask', {
            method: 'POST',
            body: formData
        });
        
        if (!response.ok) {
            throw new Error(`خطأ في الخادم: ${response.status}`);
        }
        
        const data = await response.json();
        
        // Display answer
        document.getElementById('answer').innerHTML = `<h3>✅ الرد:</h3><div>${data.answer}</div>`;
        
        // Create export links
        const questionParam = encodeURIComponent(data.question);
        const answerParam = encodeURIComponent(data.answer);
        
        document.getElementById('export').innerHTML = `
            <div class="export-links">
                <h4>📁 تصدير الرد:</h4>
                <a href="/export/pdf?question=${questionParam}&answer=${answerParam}" target="_self">📄 تحميل PDF</a>
                <a href="/export/docx?question=${questionParam}&answer=${answerParam}" target="_self">📝 تحميل Word</a>
            </div>
        `;
        
    } catch (error) {
        console.error('Error:', error);
        document.getElementById('answer').innerHTML = `<p style="color: red;">❌ حدث خطأ: ${error.message}</p>`;
    }
}

// Allow Enter key to send question
document.getElementById('query').addEventListener('keypress', function(e) {
    if (e.key === 'Enter' && e.ctrlKey) {
        sendQuestion();
    }
});
</script>
</body>
</html>
    """

@app.post("/ask")
def ask_api(query: str = Form(...)):
    try:
        answer = ask_question(query)
        return {"question": query, "answer": answer}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.get("/export/pdf")
def export_pdf(question: str = Query(...), answer: str = Query(...)):
    try:
        # Create unique filename in temp directory
        filename = f"{uuid.uuid4().hex}.pdf"
        filepath = os.path.join(TEMP_DIR, filename)
        
        # Create PDF
        pdf = FPDF()
        pdf.add_page()
        
        # Add Unicode font support for Arabic (if you have Arabic font)
        try:
            pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
            pdf.set_font('DejaVu', size=12)
        except:
            # Fallback to Arial if Arabic font not available
            pdf.set_font("Arial", size=12)
        
        # Clean HTML tags from answer
        clean_answer = re.sub('<[^<]+?>', '', answer)
        clean_question = re.sub('<[^<]+?>', '', question)
        
        # Add content
        content = f"السؤال:\n{clean_question}\n\nالإجابة:\n{clean_answer}"
        pdf.multi_cell(0, 10, content.encode('latin-1', 'replace').decode('latin-1'))
        
        # Save PDF
        pdf.output(filepath)
        
        return FileResponse(
            path=filepath,
            media_type="application/pdf",
            filename="الرد-القانوني.pdf",
            headers={
                "Content-Disposition": "attachment; filename=الرد-القانوني.pdf",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "GET",
                "Access-Control-Allow-Headers": "*"
            }
        )
        
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"خطأ في إنشاء ملف PDF: {str(e)}"})

@app.get("/export/docx")
def export_docx(question: str = Query(...), answer: str = Query(...)):
    try:
        # Create unique filename in temp directory
        filename = f"{uuid.uuid4().hex}.docx"
        filepath = os.path.join(TEMP_DIR, filename)
        
        # Create Word document
        doc = Document()
        
        # Clean HTML tags
        clean_answer = re.sub('<[^<]+?>', '', answer)
        clean_question = re.sub('<[^<]+?>', '', question)
        
        # Add content
        doc.add_heading('الرد القانوني', 0)
        doc.add_heading('السؤال:', level=1)
        doc.add_paragraph(clean_question)
        doc.add_heading('الإجابة:', level=1)
        doc.add_paragraph(clean_answer)
        
        # Save document
        doc.save(filepath)
        
        return FileResponse(
            path=filepath,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="الرد-القانوني.docx",
            headers={
                "Content-Disposition": "attachment; filename=الرد-القانوني.docx",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "GET",
                "Access-Control-Allow-Headers": "*"
            }
        )
        
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"خطأ في إنشاء ملف Word: {str(e)}"})

# Simple cleanup function
def cleanup_old_files():
    """Clean up old files"""
    try:
        if os.path.exists(TEMP_DIR):
            current_time = time.time()
            for filename in os.listdir(TEMP_DIR):
                filepath = os.path.join(TEMP_DIR, filename)
                if os.path.isfile(filepath):
                    # Delete files older than 1 hour
                    if current_time - os.path.getmtime(filepath) > 3600:
                        os.remove(filepath)
                        print(f"🗑️ Cleaned up old file: {filepath}")
    except Exception as e:
        print(f"❌ Error during cleanup: {e}")

# Clean up on startup
@app.on_event("startup")
async def startup_event():
    """Clean up old temp files on startup"""
    cleanup_old_files()