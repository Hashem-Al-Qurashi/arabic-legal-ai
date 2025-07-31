# COMPLETE CLEAN VERSION - Replace your entire backend/app/api/export.py

"""
Export Router - Clean Debug Version with Perfect Copy Formatting
"""
from fastapi import APIRouter, Query, HTTPException, Depends
from fastapi.responses import StreamingResponse
from datetime import datetime
import io
import re
import traceback
from typing import Optional
import urllib.parse

# Try multiple import paths for dependencies
print("🔍 Testing import paths...")

try:
    from app.dependencies.simple_auth import get_optional_current_user
    print("✅ Import 1: app.dependencies.simple_auth.get_optional_current_user")
except ImportError as e:
    print(f"❌ Import 1 failed: {e}")
    try:
        from app.dependencies.auth import get_current_user_optional as get_optional_current_user
        print("✅ Import 2: app.dependencies.auth.get_current_user_optional")
    except ImportError as e2:
        print(f"❌ Import 2 failed: {e2}")
        # Create a dummy function if all imports fail
        def get_optional_current_user():
            return None
        print("⚠️ Using dummy auth function")

try:
    from app.models.user import User
    print("✅ Import: app.models.user.User")
except ImportError as e:
    print(f"❌ User model import failed: {e}")
    # Create dummy User class
    class User:
        def __init__(self):
            self.full_name = "Unknown User"
    print("⚠️ Using dummy User class")

router = APIRouter(tags=["Export"])

def apply_copy_formatting_to_word(content: str) -> str:
    """
    Apply the same clean formatting logic as the copy function
    Handles BOTH HTML tags AND markdown syntax - CLEAN VERSION
    """
    print("🎨 Applying copy formatting...")
    
    try:
        clean_content = content
        
        # Convert HTML headings to structured text
        clean_content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'\n━━━ \1 ━━━\n\n', clean_content, flags=re.IGNORECASE)
        clean_content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'\n▎ \1\n\n', clean_content, flags=re.IGNORECASE) 
        clean_content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'\n▸ \1\n\n', clean_content, flags=re.IGNORECASE)
        
        # Convert MARKDOWN headings to structured text
        clean_content = re.sub(r'^####\s*(.*?)$', r'\n▸ \1\n\n', clean_content, flags=re.MULTILINE)
        clean_content = re.sub(r'^###\s*(.*?)$', r'\n▸ \1\n\n', clean_content, flags=re.MULTILINE)
        clean_content = re.sub(r'^##\s*(.*?)$', r'\n▎ \1\n\n', clean_content, flags=re.MULTILINE)
        clean_content = re.sub(r'^#\s*(.*?)$', r'\n━━━ \1 ━━━\n\n', clean_content, flags=re.MULTILINE)
        
        # Convert legal-point divs
        clean_content = re.sub(
            r'<div class="legal-point"><strong>(.*?)</strong><p>(.*?)</p></div>',
            r'\1 \2\n\n',
            clean_content,
            flags=re.IGNORECASE
        )
        
        # CRITICAL: Convert MARKDOWN bold to clean text (NO ASTERISKS!)
        clean_content = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_content)
        clean_content = re.sub(r'\*(.*?)\*', r'\1', clean_content)
        
        # Convert HTML bold/strong to clean text
        clean_content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'\1', clean_content, flags=re.IGNORECASE)
        clean_content = re.sub(r'<b[^>]*>(.*?)</b>', r'\1', clean_content, flags=re.IGNORECASE)
        
        # Convert emphasis to clean text  
        clean_content = re.sub(r'<em[^>]*>(.*?)</em>', r'\1', clean_content, flags=re.IGNORECASE)
        clean_content = re.sub(r'<i[^>]*>(.*?)</i>', r'\1', clean_content, flags=re.IGNORECASE)
        
        # Convert lists to clean bullet points
        clean_content = re.sub(r'<ul[^>]*>', '\n', clean_content, flags=re.IGNORECASE)
        clean_content = re.sub(r'</ul>', '\n', clean_content, flags=re.IGNORECASE)
        clean_content = re.sub(r'<ol[^>]*>', '\n', clean_content, flags=re.IGNORECASE) 
        clean_content = re.sub(r'</ol>', '\n', clean_content, flags=re.IGNORECASE)
        clean_content = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', clean_content, flags=re.IGNORECASE)
        
        # Convert paragraphs with proper spacing
        clean_content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', clean_content, flags=re.IGNORECASE)
        
        # Convert line breaks
        clean_content = re.sub(r'<br\s*/?>', '\n', clean_content, flags=re.IGNORECASE)
        
        # Convert divs to clean text
        clean_content = re.sub(r'<div[^>]*>(.*?)</div>', r'\1\n', clean_content, flags=re.IGNORECASE)
        
        # Remove any remaining HTML tags
        clean_content = re.sub(r'<[^>]*>', '', clean_content)
        
        # Convert HTML entities
        clean_content = clean_content.replace('&nbsp;', ' ')
        clean_content = clean_content.replace('&amp;', '&')
        clean_content = clean_content.replace('&lt;', '<')
        clean_content = clean_content.replace('&gt;', '>')
        clean_content = clean_content.replace('&quot;', '"')
        clean_content = clean_content.replace('&#39;', "'")
        clean_content = clean_content.replace('&hellip;', '...')
        
        # Clean up whitespace while preserving structure
        clean_content = re.sub(r'[ \t]+', ' ', clean_content)
        clean_content = re.sub(r'\n[ \t]+', '\n', clean_content)
        clean_content = re.sub(r'[ \t]+\n', '\n', clean_content)
        clean_content = re.sub(r'\n{4,}', '\n\n\n', clean_content)
        
        # Ensure proper Arabic text flow
        clean_content = re.sub(r'([أ-ي])\n+(أولاً|ثانياً|ثالثاً|رابعاً|خامساً)', r'\1\n\n\2', clean_content)
        clean_content = re.sub(r'([أ-ي])\n+(\d+\.)', r'\1\n\n\2', clean_content)
        clean_content = re.sub(r'([أ-ي])\n+(▸|•)', r'\1\n\n\2', clean_content)
        
        # Final structure enhancement for readability
        clean_content = re.sub(r'(━━━.*━━━)\n{1,2}([^▸•])', r'\1\n\n\2', clean_content)
        clean_content = re.sub(r'(▸.*?)\n{1,2}([^▸•▎])', r'\1\n\n\2', clean_content)
        
        # Clean bullet formatting
        clean_content = re.sub(r'•\s*', '• ', clean_content)
        clean_content = re.sub(r'▸\s*', '▸ ', clean_content)
        
        # Remove excess whitespace at beginning/end
        clean_content = re.sub(r'\n{2,}$', '\n', clean_content)
        clean_content = re.sub(r'^\n{2,}', '', clean_content)
        
        result = clean_content.strip()
        print(f"✅ Copy formatting applied. Length: {len(result)} chars")
        return result
        
    except Exception as e:
        print(f"❌ Copy formatting failed: {e}")
        # Fallback to basic cleaning
        basic_clean = re.sub(r'<[^>]*>', '', content)
        return basic_clean.replace('&nbsp;', ' ').strip()

def create_formatted_docx_stream(question: str, answer: str):
    """
    Create Word document with beautiful formatting
    """
    print("📝 Starting formatted DOCX creation...")
    buffer = io.BytesIO()
    
    try:
        print("📦 Importing python-docx...")
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        print("✅ python-docx imported successfully")
        
        print("📄 Creating document...")
        doc = Document()
        
        print("🧹 Cleaning content with copy formatting...")
        clean_question = apply_copy_formatting_to_word(question)
        clean_answer = apply_copy_formatting_to_word(answer)
        
        print(f"📝 Question length after cleaning: {len(clean_question)}")
        print(f"📝 Answer length after cleaning: {len(clean_answer)}")
        
        print("📋 Adding title...")
        title = doc.add_heading('المساعد القانوني الذكي 🇸🇦', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph('استشارة قانونية ذكية مبنية على القانون السعودي')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph('━' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Empty line
        
        print("❓ Adding question with structure...")
        doc.add_heading('📋 السؤال:', level=1)
        
        # Process question lines
        if clean_question:
            question_lines = clean_question.split('\n')
            for line in question_lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith('━━━') and line.endswith('━━━'):
                    # Main heading
                    heading_text = line.replace('━━━', '').strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=2)
                elif line.startswith('▎'):
                    # Sub heading
                    heading_text = line.replace('▎', '').strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=3)
                elif line.startswith('▸'):
                    # Minor heading
                    heading_text = line.replace('▸', '').strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=4)
                elif line.startswith('•'):
                    # Bullet point
                    doc.add_paragraph(line, style='List Bullet')
                else:
                    # Regular paragraph
                    if len(line) > 5:  # Avoid very short lines
                        doc.add_paragraph(line)
        
        doc.add_paragraph()  # Empty line between sections
        
        print("✅ Adding answer with structure...")
        doc.add_heading('📝 الإجابة:', level=1)
        
        # Process answer lines
        if clean_answer:
            answer_lines = clean_answer.split('\n')
            for line in answer_lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith('━━━') and line.endswith('━━━'):
                    # Main heading
                    heading_text = line.replace('━━━', '').strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=2)
                elif line.startswith('▎'):
                    # Sub heading
                    heading_text = line.replace('▎', '').strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=3)
                elif line.startswith('▸'):
                    # Minor heading
                    heading_text = line.replace('▸', '').strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=4)
                elif line.startswith('•'):
                    # Bullet point
                    doc.add_paragraph(line, style='List Bullet')
                else:
                    # Regular paragraph
                    if len(line) > 5:  # Avoid very short lines
                        doc.add_paragraph(line)
        
        # Add footer
        doc.add_paragraph()
        doc.add_paragraph()
        footer_line = doc.add_paragraph('━' * 50)
        footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp_text = f"تاريخ الإنشاء: {datetime.now().strftime('%Y-%m-%d الساعة %H:%M')}"
        footer_para = doc.add_paragraph(timestamp_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer_run = disclaimer.add_run(
            'تنبيه: هذه الاستشارة القانونية مبنية على الذكاء الاصطناعي وتهدف للإرشاد العام. '
            'للحصول على استشارة قانونية دقيقة، يُنصح بالتواصل مع محامٍ مختص.'
        )
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.italic = True
        
        print("💾 Saving to buffer...")
        doc.save(buffer)
        print("✅ Document saved successfully")
        
    except ImportError as e:
        print(f"❌ python-docx not installed: {e}")
        raise HTTPException(status_code=500, detail=f"python-docx library not installed: {str(e)}")
    except Exception as e:
        print(f"❌ DOCX creation failed: {e}")
        print(f"❌ Full traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"DOCX creation error: {str(e)}")
    
    buffer.seek(0)
    print(f"📊 Buffer size: {len(buffer.getvalue())} bytes")
    return buffer

@router.get("/docx")
async def export_docx_debug(
    question: str = Query(..., description="The legal question asked"),
    answer: str = Query(..., description="The AI assistant's response"),
    current_user: Optional[User] = Depends(get_optional_current_user)
):
    """
    DEBUG VERSION - Export with detailed logging and copy formatting
    """
    print("🚀 =================================")
    print("🚀 EXPORT DOCX DEBUG - STARTING")
    print("🚀 =================================")
    
    try:
        print(f"📝 Question received: {len(question)} characters")
        print(f"📝 Answer received: {len(answer)} characters")
        print(f"👤 User: {current_user.full_name if current_user and hasattr(current_user, 'full_name') else 'Guest'}")
        
        # SAFETY: Limit content length to avoid issues
        if len(question) > 5000:
            question = question[:5000] + "... (محتوى مقطوع)"
            print("⚠️ Question truncated to 5000 chars")
        
        if len(answer) > 10000:
            answer = answer[:10000] + "... (محتوى مقطوع)"
            print("⚠️ Answer truncated to 10000 chars")
        
        # Log first 100 chars of each
        print(f"📝 Question preview: {question[:100]}...")
        print(f"📝 Answer preview: {answer[:100]}...")
        
        print("🏗️ Creating formatted DOCX buffer...")
        docx_buffer = create_formatted_docx_stream(question, answer)
        
        print("📁 Generating filename...")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Arabic filename with proper encoding
        arabic_filename = f"استشارة_قانونية_منسقة_{timestamp}.docx"
        
        # URL-encode the Arabic filename for HTTP headers
        encoded_filename = urllib.parse.quote(arabic_filename.encode('utf-8'))
        print(f"📁 Filename: {arabic_filename} (encoded: {encoded_filename})")
        
        # Use RFC 5987 format for international filenames
        filename_header = f"filename*=UTF-8''{encoded_filename}"
        
        print("📤 Preparing response...")
        buffer_content = docx_buffer.read()
        print(f"📊 Final buffer size: {len(buffer_content)} bytes")
        
        if len(buffer_content) == 0:
            raise HTTPException(status_code=500, detail="Generated DOCX file is empty")
        
        print("✅ =================================")
        print("✅ EXPORT DOCX DEBUG - SUCCESS")
        print("✅ =================================")
        
        return StreamingResponse(
            io.BytesIO(buffer_content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; {filename_header}",
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "Cache-Control": "no-cache"
            }
        )
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print("❌ =================================")
        print("❌ EXPORT DOCX DEBUG - FAILED")
        print("❌ =================================")
        print(f"❌ Error: {str(e)}")
        print(f"❌ Error type: {type(e).__name__}")
        print(f"❌ Full traceback:")
        print(traceback.format_exc())
        print("❌ =================================")
        
        raise HTTPException(
            status_code=500, 
            detail=f"Export failed: {str(e)}"
        )

@router.get("/test")
async def test_export():
    """Simple test endpoint"""
    print("🧪 Test endpoint called")
    
    # Test python-docx import
    try:
        from docx import Document
        docx_status = "✅ Available"
    except ImportError as e:
        docx_status = f"❌ Not available: {e}"
    
    # Test auth import
    try:
        from app.dependencies.simple_auth import get_optional_current_user
        auth_status = "✅ simple_auth available"
    except ImportError:
        try:
            from app.dependencies.auth import get_current_user_optional
            auth_status = "✅ auth available"
        except ImportError as e:
            auth_status = f"❌ No auth available: {e}"
    
    return {
        "status": "Export router is working",
        "timestamp": datetime.now().isoformat(),
        "python_docx": docx_status,
        "auth_dependency": auth_status,
        "available_endpoints": ["/docx", "/test"]
    }

print("🎯 Export router loaded with DEBUG version")