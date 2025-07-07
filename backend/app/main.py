"""
Ultimate Arabic Legal Assistant - Complete Feature Set
Combines JWT authentication, guest access, chat system, and direct RAG
"""
from app.api.chat import router as chat_router
from fastapi import FastAPI, HTTPException, Query, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from datetime import datetime
import io
import os
import re

# Import routers
from app.api.simple_auth import router as auth_router  # ← CORRECT!

# Initialize database tables
from app.database import engine, Base
Base.metadata.create_all(bind=engine)
print("✅ Database tables created!")

# ✅ CORS using environment variable with fallback
def get_cors_origins():
    """Get CORS origins from environment or use defaults"""
    cors_origins = os.getenv("CORS_ORIGINS")
    if cors_origins:
        return [origin.strip() for origin in cors_origins.split(",")]
    return ["http://localhost:3000", "http://127.0.0.1:3000"]

# Create FastAPI application
app = FastAPI(
    title="Arabic Legal AI Assistant 🇸🇦",
    description="""
    🇸🇦 استشارة قانونية ومالية وإدارية ذكية مبنية على القانون السعودي
    
    Complete Arabic Legal Assistant with:
    - JWT-based user authentication and management
    - Guest access for basic consultations
    - AI-powered legal consultations with conversation memory
    - Document export capabilities with perfect Arabic support
    - Multi-domain consultation support (Legal, Financial, Administrative)
    """,
    version="2.0.0"
)

# ✅ Add CORS middleware using environment variable
app.add_middleware(
    CORSMiddleware,
    allow_origins=get_cors_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Debug log CORS configuration
cors_origins = get_cors_origins()
print(f"🌐 CORS Origins configured: {cors_origins}")

# ✅ Include authentication and user management routers
app.include_router(auth_router, prefix="/api/auth", tags=["Authentication"])
app.include_router(chat_router, prefix="/api", tags=["Chat System"])

# ✅ GUEST ACCESS: Direct RAG endpoint for non-authenticated users
# ✅ Enhanced DOCX export with perfect Arabic support
def create_enhanced_docx_stream(question: str, answer: str) -> io.BytesIO:
    """Create enhanced DOCX with perfect Arabic support"""
    buffer = io.BytesIO()
    
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create Word document
        doc = Document()
        
        # Clean HTML tags from content
        clean_question = re.sub('<[^<]+?>', '', question)
        clean_answer = re.sub('<[^<]+?>', '', answer)
        clean_answer = clean_answer.replace('&nbsp;', ' ')
        
        # Document title
        title = doc.add_heading('المساعد القانوني الذكي 🇸🇦', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('استشارة قانونية ذكية مبنية على القانون السعودي')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator line
        doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Question section
        question_heading = doc.add_heading('📋 السؤال:', level=1)
        question_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        question_para = doc.add_paragraph(clean_question)
        question_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in question_para.runs:
            run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Answer section
        answer_heading = doc.add_heading('✅ الإجابة:', level=1)
        answer_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        answer_para = doc.add_paragraph(clean_answer)
        answer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in answer_para.runs:
            run.font.size = Pt(11)
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph()
        
        footer_line = doc.add_paragraph('─' * 50)
        footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp_text = f"تاريخ الإنشاء: {datetime.now().strftime('%Y-%m-%d الساعة %H:%M')}"
        footer_para = doc.add_paragraph(timestamp_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer_run = disclaimer.add_run(
            'تنبيه: هذه الاستشارة القانونية مبنية على الذكاء الاصطناعي وتهدف للإرشاد العام. '
            'للحصول على استشارة قانونية دقيقة، يُنصح بالتواصل مع محامٍ مختص.'
        )
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.italic = True
        
        # Save to buffer
        doc.save(buffer)
        
    except Exception as e:
        print(f"DOCX creation error: {e}")
        # Create simple fallback document
        from docx import Document
        doc = Document()
        doc.add_heading('خطأ في إنشاء المستند', 0)
        doc.add_paragraph('حدث خطأ أثناء إنشاء المستند. يرجى المحاولة مرة أخرى.')
        doc.add_paragraph(f'تفاصيل الخطأ: {str(e)}')
        doc.save(buffer)
    
    buffer.seek(0)
    return buffer

@app.get("/export/docx")
async def export_docx(question: str = Query(...), answer: str = Query(...)):
    """Export legal response as DOCX with perfect Arabic support"""
    try:
        print(f"📝 Generating enhanced DOCX export...")
        
        # Create enhanced DOCX
        docx_buffer = create_enhanced_docx_stream(question, answer)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"legal_consultation_{timestamp}.docx"
        
        print(f"✅ Enhanced DOCX export generated: {filename}")
        
        return StreamingResponse(
            io.BytesIO(docx_buffer.read()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        )
        
    except Exception as e:
        print(f"❌ DOCX export error: {e}")
        raise HTTPException(status_code=500, detail=f"خطأ في إنشاء ملف Word: {str(e)}")

# ✅ Legal categories endpoint
@app.get("/api/categories")
async def get_legal_categories():
    """Get available legal categories for consultation classification"""
    categories = [
        {"id": "commercial", "name": "القانون التجاري", "emoji": "💼"},
        {"id": "labor", "name": "قانون العمل", "emoji": "👷"},
        {"id": "real_estate", "name": "القانون العقاري", "emoji": "🏠"},
        {"id": "family", "name": "الأحوال الشخصية", "emoji": "👨‍👩‍👧‍👦"},
        {"id": "criminal", "name": "القانون الجنائي", "emoji": "⚖️"},
        {"id": "administrative", "name": "القانون الإداري", "emoji": "🏛️"},
        {"id": "general", "name": "استشارة عامة", "emoji": "📋"}
    ]
    return {"categories": categories}

# ✅ Enhanced health check endpoint
@app.get("/")
async def root():
    """Enhanced API health check with full feature overview"""
    return {
        "status": "healthy",
        "message": "🇸🇦 Arabic Legal AI Assistant - Ultimate Edition",
        "version": "2.0.0",
        "features": {
            "authentication": "JWT-based user management",
            "guest_access": "Direct legal consultation for guests",
            "chat_system": "Conversation memory for authenticated users",
            "export": "Enhanced DOCX with perfect Arabic support",
            "categories": "Multi-domain legal classification",
            "ai_engine": "DeepSeek-powered RAG system"
        },
        "endpoints": {
            "guest": "/api/ask",
            "auth": "/api/auth/*",
            "chat": "/api/chat/*",
            "users": "/api/users/*",
            "export": "/export/docx",
            "categories": "/api/categories"
        },
        "cors_origins": get_cors_origins(),
        "environment": os.getenv("ENVIRONMENT", "development"),
        "timestamp": datetime.now().isoformat()
    }

@app.get("/health")
async def health_check():
    """Simple health check endpoint"""
    return {
        "status": "healthy", 
        "timestamp": datetime.now().isoformat(),
        "service": "Arabic Legal Assistant - Ultimate Edition",
        "features": ["jwt_auth", "guest_access", "chat_memory", "docx_export", "arabic_support"]
    }

print("🚀 Arabic Legal Assistant Ultimate Edition started!")
print("✅ Features: JWT Auth + Guest Access + Chat Memory + Enhanced Export")
print(f"🌐 CORS configured for: {get_cors_origins()}")