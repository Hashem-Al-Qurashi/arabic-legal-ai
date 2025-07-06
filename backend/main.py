from fastapi import FastAPI, Form, Query, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from rag_engine import ask_question
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from datetime import datetime

app = FastAPI(
    title="Arabic Legal AI Assistant 🇸🇦",
    description="استشارة قانونية ومالية وإدارية ذكية مبنية على القانون السعودي.",
    version="1.0.0"
)

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def create_docx_stream(question: str, answer: str) -> io.BytesIO:
    """Create enhanced DOCX with perfect Arabic support"""
    buffer = io.BytesIO()
    
    try:
        # Create Word document
        doc = Document()
        
        # Clean HTML tags from content
        clean_question = re.sub('<[^<]+?>', '', question)
        clean_answer = re.sub('<[^<]+?>', '', answer)
        
        # Document title
        title = doc.add_heading('المساعد القانوني الذكي 🇸🇦', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('استشارة قانونية ذكية مبنية على القانون السعودي')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator line
        doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        doc.add_paragraph()
        
        # Question section
        question_heading = doc.add_heading('📋 السؤال:', level=1)
        question_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        question_para = doc.add_paragraph(clean_question)
        question_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Make question text slightly larger
        for run in question_para.runs:
            run.font.size = Pt(12)
        
        # Add space between sections
        doc.add_paragraph()
        
        # Answer section
        answer_heading = doc.add_heading('✅ الإجابة:', level=1)
        answer_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        answer_para = doc.add_paragraph(clean_answer)
        answer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Make answer text readable
        for run in answer_para.runs:
            run.font.size = Pt(11)
        
        # Add space before footer
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Footer with timestamp
        footer_line = doc.add_paragraph('─' * 50)
        footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp_text = f"تاريخ الإنشاء: {datetime.now().strftime('%Y-%m-%d الساعة %H:%M')}"
        footer_para = doc.add_paragraph(timestamp_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer_run = disclaimer.add_run('تنبيه: هذه الاستشارة القانونية مبنية على الذكاء الاصطناعي وتهدف للإرشاد العام. '
                                          'للحصول على استشارة قانونية دقيقة، يُنصح بالتواصل مع محامٍ مختص.')
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.italic = True
        
        # Save to buffer
        doc.save(buffer)
        
    except Exception as e:
        print(f"DOCX creation error: {e}")
        # Create simple fallback document
        doc = Document()
        doc.add_heading('خطأ في إنشاء المستند', 0)
        doc.add_paragraph('حدث خطأ أثناء إنشاء المستند. يرجى المحاولة مرة أخرى.')
        doc.add_paragraph(f'تفاصيل الخطأ: {str(e)}')
        doc.save(buffer)
    
    buffer.seek(0)
    return buffer

@app.post("/ask")
async def ask_api(query: str = Form(...)):
    """Process legal question and return answer"""
    try:
        if not query or not query.strip():
            raise HTTPException(status_code=400, detail="السؤال مطلوب")
        
        print(f"🤖 Processing question: {query[:50]}...")
        answer = ask_question(query.strip())
        
        return {
            "question": query.strip(),
            "answer": answer,
            "timestamp": datetime.now().isoformat(),
            "status": "success"
        }
    except Exception as e:
        print(f"❌ Error processing question: {e}")
        raise HTTPException(status_code=500, detail="حدث خطأ في معالجة السؤال")

@app.get("/export/docx")
async def export_docx(question: str = Query(...), answer: str = Query(...)):
    """Export legal response as DOCX with perfect Arabic support"""
    try:
        print(f"📝 Generating DOCX export...")
        
        # Create DOCX in memory
        docx_buffer = create_docx_stream(question, answer)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"legal_response_{timestamp}.docx"
        
        print(f"✅ DOCX export generated successfully")
        
        # Return as streaming response
        return StreamingResponse(
            io.BytesIO(docx_buffer.read()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        )
        
    except Exception as e:
        print(f"❌ DOCX export error: {e}")
        raise HTTPException(status_code=500, detail=f"خطأ في إنشاء ملف Word: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy", 
        "timestamp": datetime.now().isoformat(),
        "service": "Arabic Legal Assistant",
        "features": ["legal_consultation", "docx_export", "arabic_support"]
    }

# Optional: Add endpoint to get available legal categories
@app.get("/categories")
async def get_legal_categories():
    """Get available legal categories"""
    categories = [
        {"id": "commercial", "name": "القانون التجاري", "emoji": "💼"},
        {"id": "labor", "name": "قانون العمل", "emoji": "👷"},
        {"id": "real_estate", "name": "القانون العقاري", "emoji": "🏠"},
        {"id": "family", "name": "الأحوال الشخصية", "emoji": "👨‍👩‍👧‍👦"},
        {"id": "criminal", "name": "القانون الجنائي", "emoji": "⚖️"},
        {"id": "administrative", "name": "القانون الإداري", "emoji": "🏛️"},
        {"id": "general", "name": "استشارة عامة", "emoji": "📋"}
    ]
    return {"categories": categories}

if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting Arabic Legal Assistant API...")
    print("📝 DOCX Export: Enabled with perfect Arabic support")
    print("📄 PDF Export: Disabled (Arabic text issues)")
    uvicorn.run(app, host="0.0.0.0", port=8000)